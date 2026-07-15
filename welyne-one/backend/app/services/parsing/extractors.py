"""
Extraction texte brute — PyMuPDF / python-docx / Docling, avec repli OCR Tesseract
(fra+eng+ara) quand la couche texte est trop pauvre (< 200 caractères/page).
"""
from __future__ import annotations

import io
import logging
from pathlib import Path

logger = logging.getLogger("welyne.a3.extractors")

OCR_THRESHOLD_CHARS_PER_PAGE = 200


def extract_text(file_path: str, mime: str) -> tuple[str, bool]:
    """Retourne (texte_brut, ocr_used)."""
    path = Path(file_path)

    if mime == "application/pdf" or path.suffix.lower() == ".pdf":
        return _extract_pdf(path)

    if mime in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ) or path.suffix.lower() == ".docx":
        return _extract_docx(path), False

    if mime.startswith("image/"):
        return _ocr_image(path), True

    if mime == "text/plain" or path.suffix.lower() == ".txt":
        # A2 (§6-A2) : profil LinkedIn collé/exporté en texte brut par le
        # recruteur (pas de PDF) — pas d'extraction ni d'OCR nécessaire.
        return path.read_text(encoding="utf-8", errors="replace"), False

    raise ValueError(f"Type de fichier non supporté pour l'extraction : {mime} ({path.suffix})")


def _extract_pdf(path: Path) -> tuple[str, bool]:
    import fitz  # PyMuPDF

    doc = fitz.open(str(path))
    pages_text = []
    for page in doc:
        pages_text.append(page.get_text())
    text = "\n".join(pages_text)

    avg_chars_per_page = len(text) / max(len(doc), 1)
    if avg_chars_per_page < OCR_THRESHOLD_CHARS_PER_PAGE:
        logger.info("Couche texte faible (%.0f car/page) — repli OCR pour %s", avg_chars_per_page, path.name)
        ocr_text = _ocr_pdf(doc)
        return ocr_text, True

    return text, False


def _ocr_pdf(doc) -> str:
    import pytesseract
    from PIL import Image

    texts = []
    for page in doc:
        pix = page.get_pixmap(dpi=200)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        texts.append(pytesseract.image_to_string(img, lang="fra+eng+ara"))
    return "\n".join(texts)


def _extract_docx(path: Path) -> str:
    import docx

    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _ocr_image(path: Path) -> str:
    import pytesseract
    from PIL import Image

    return pytesseract.image_to_string(Image.open(path), lang="fra+eng+ara")


def detect_language(text: str) -> str:
    """fr | en | ar — repli 'fr' si indétectable (CV tunisiens majoritairement FR)."""
    try:
        from langdetect import detect

        code = detect(text[:2000])
        return code if code in ("fr", "en", "ar") else "fr"
    except Exception:  # noqa: BLE001
        return "fr"